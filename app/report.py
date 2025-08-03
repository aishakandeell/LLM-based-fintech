import os
import pandas as pd
from docx import Document
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
import cohere


COHERE_API_KEY = os.getenv("CO_API_KEY") or "your_fallback_api_key"
co = cohere.Client(COHERE_API_KEY)

def clean_value(val):
    val = str(val).strip()
    if val.lower() in ["nan", "–", "-", "", "none"]:
        return None
    return val

def generate_summary_overview(analysis_context: dict, company_name: str, industry_name: str, output_dir: str, custom_filename: str = None) -> str:


    doc = Document()

    # Build concise context summary
    context_summary = f"""
    EXECUTIVE SUMMARY:
    {analysis_context['executive_summary']}

    COMPANY OVERVIEW:
    {analysis_context['company_overview']}

    INDUSTRY ANALYSIS:
    {analysis_context['industry_analysis']}

    FINANCIAL ANALYSIS:
    {analysis_context['financial_analysis']}

    SWOT ANALYSIS:
    {analysis_context['swot_analysis']}

    STRATEGIC INITIATIVES AND RECOMMENDATIONS:
    {analysis_context['strategic_initiatives']}
    """

    # Define overview prompt
    overview_prompt = f"""
    You are a senior business consultant. Using the following detailed company analysis, produce a concise Overview Report for {company_name} in the {industry_name} industry.

    The report should include:

    1. Explanation: Summarize the company's current situation, key strengths, weaknesses, and overall position.

    2. Comparison to Macro State: Compare the company's position to the broader industry/economic context in Egypt and globally, highlighting where it stands out or lags.

    3. Suggestions/Recommendations for Future Steps: Provide actionable strategic recommendations categorized clearly (Exit, Monitor, Restructure, Invest, Strategic Adjustments), focusing on top priorities only.

    Formatting:

    - Professional business report style
    - Maximum of 2 pages worth of content (approx 600-800 words)
    - Structured with clear section headings and short bullet points where needed
    - Concise and to the point for executive reference

    Here is the full company analysis:

    {context_summary}
    """

    # Generate using Cohere
    overview_response = co.chat(
        model="command-nightly",
        message=overview_prompt,
        temperature=0.4
    ).text.strip()

    # Write to Word document
    doc.add_heading("Company Summary Overview", level=1)
    doc.add_paragraph(overview_response)

    # Save overview report
    filename = f"{custom_filename}.docx" if custom_filename else f"summary_overview_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)

    return output_path


def generate_report(df: pd.DataFrame, company_name: str = None, industry_name: str=None, custom_filename: str = None, overview_filename: str = None) -> tuple[str, str]:

    doc = Document()
    analysis_context = {}

    summary_response = co.chat(
        model="command-nightly",
        message=(
            f"Provide a concise Executive Summary for the company '{company_name}', located in Egypt. "
            "Include a short introduction, current financial or operational highlights, "
            "and recent strategic developments if available. "
            "Write in a business report style, under 150 words."
            "Do NOT use placeholder words like [ai_recommendations]; only use the company's actual name."
        ),
        temperature=0.5
        ).text.strip()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(summary_response)
    analysis_context['executive_summary'] = summary_response


    # AI prompt to generate detailed company overview
    company_overview = co.chat(
        model="command-nightly",
        message=(
            f"Provide a detailed Company Overview for '{company_name}', located in Egypt, in formal business report style. "
            "Include the following structured sections:\n\n"
            "1. Company Overview\n"
            "- A general introductory paragraph about the company.\n\n"
            "1.1 Historical Context\n"
            "- When it was established.\n"
            "- Key facilities with short bullet points including location, capacity, and purpose.\n\n"
            "1.2 Current Status\n"
            "- Recent operational or financial status.\n"
            "- Key challenges or recent strategic partnerships, with bullet points summarizing main issues.\n\n"
            "Return all text formatted for inclusion in a Word business report."
        ),
        temperature=0.5
    ).text.strip()
    # Add Company Overview section to report
    doc.add_heading("1. Company Overview", level=1)
    doc.add_paragraph(company_overview)
    analysis_context['company_overview'] = company_overview


    industry_analysis = co.chat(
        model="command-nightly",
        message=(
            f"Provide an Industry and Market Analysis for '{industry_name}' in Egypt. "
            "Write in a structured business report format with the following sections:\n\n"
            "2. Industry and Market Analysis\n\n"
            "2.1 Egyptian {industry_name} Market Overview\n"
            "- A paragraph summarizing the market size, growth rate, and major drivers.\n"
            "- Bullet points listing key drivers or initiatives.\n\n"
            "2.2 Egypt's Position in the Global {industry_name} Market\n"
            "- A paragraph summarizing Egypt's global ranking and role.\n"
            "- Bullet points on export performance, market reach, European market position, and growth trajectory.\n\n"
            "Format exactly like a business report with proper headings."
        ),
        temperature=0.5
    ).text.strip()
    doc.add_heading("2. Industry and Market Analysis", level=1)
    doc.add_paragraph(industry_analysis)
    analysis_context['industry_analysis'] = industry_analysis


    financial_analysis_prompt = f"""
        You are a financial analyst. Using the following extracted financial table, produce a detailed Financial Analysis section
        in formal business report format with the following subsections and formatting:

        3. Financial Analysis

        3.1 Historical Financial Performance

        - Revenue Trends:
            • Describe overall revenue trends with CAGR calculated and revenue changes year by year in bullet points.

        - Profitability Metrics:
            • Provide gross profit margins, EBITDA margins, net profits, and net losses as bullet points.

        - Balance Sheet Position:
            • Describe changes in total assets, liabilities, equity, and debt-to-assets ratio as bullet points.

        - Cash Flow Analysis:
            • Summarize cash flows and working capital trends as bullet points.

        - Financial Efficiency Ratios:
            • Return on Assets, Return on Equity and any efficiency ratios available, as bullet points.

        3.2 Current Financial Challenges
        - Based on the above analysis, list bullet points on financial challenges similar to:
            • Severely Constrained Revenue Generation
            • Negative Operational Profitability
            • High Debt Burden
            • Technical Insolvency
            • Negative Operating Cash Flow
            • Capital Requirements

        3.3 Comparison with Industry Benchmarks
        - Compare the company’s key metrics (gross margin, EBITDA margin, net profit margin, ROA, debt-to-assets ratio) with average industry benchmarks for Egypt (search and include benchmark values with sources/links inline).

    Return the entire section formatted clearly with all subheadings and bullet points, in formal business report style.
    Here is the data table:

    {df.to_markdown(index=False)}
    """
    financial_analysis = co.chat(
        model="command-nightly",
        message=financial_analysis_prompt,
        temperature=0.3
    ).text.strip()

    doc.add_heading("3. Financial Analysis", level=1)
    doc.add_paragraph(financial_analysis)
    analysis_context['financial_analysis'] = financial_analysis


    swot_prompt = f"""
    You are a strategic consultant. Using the company's financials and strategic context below, produce Section 4. SWOT Analysis in a business report format with the following structure and formatting:

    4. SWOT Analysis

    4.1 Strengths
    List the top 5 strengths with:
    - The **Strength** as a bold bullet point title.
    - Under each strength, add a short **Description** explaining why it is a strength.
    - Then add **Strategic Implications** as further indented bullet points, explaining how this strength can be leveraged for strategic advantage.

    4.2 Weaknesses
    List the top 5 weaknesses with:
    - The **Weakness** as a bold bullet point title.
    - Under each weakness, add a short **Description** explaining why it is a weakness.
    - Then add **Strategic Implications** as further indented bullet points, explaining its risks or what must be addressed.

    Here is the company's data for context:

    {df.to_markdown(index=False)}

    Return the output in clear structured bullets, no table, following this exact style:

    • **Strength Title**
        Description: xxxxxx
        Strategic Implications:
            - xxxxx
            - xxxxx

    • **Weakness Title**
        Description: xxxxxx
        Strategic Implications:
            - xxxxx
            - xxxxx
    """
    swot_analysis = co.chat(
        model="command-nightly",
        message=swot_prompt,
        temperature=0.5
    ).text.strip()
    doc.add_heading("4. SWOT Analysis", level=1)
    doc.add_paragraph(swot_analysis)
    analysis_context['swot_analysis'] = swot_analysis

    context_summary = f"""
        EXECUTIVE SUMMARY:
        {analysis_context['executive_summary']}

        COMPANY OVERVIEW:
        {analysis_context['company_overview']}

        INDUSTRY ANALYSIS:
        {analysis_context['industry_analysis']}

        FINANCIAL ANALYSIS:
        {analysis_context['financial_analysis']}

        SWOT ANALYSIS:
        {analysis_context['swot_analysis']}
        """
    strategic_initiatives_response = co.chat(
        
        model="command-nightly",
        message=(
            f"Provide Strategic Initiatives and Recommendations for {company_name} in the {industry_name} industry based on the following company analysis:\n\n{context_summary}"

            "Organize the response under the following detailed structure exactly as written:\n\n"
        
            "5. Strategic Initiatives and Recommendations\n\n"
        
            "5.1 Operational Improvements\n"
            "  5.1.1 Production Efficiency Enhancement\n"
            "    - Include 2-4 initiatives with bolded titles, clear descriptions, and quantified targets where relevant.\n"
            "  5.1.2 Supply Chain Optimization\n"
            "    - Include 2-4 initiatives with bolded titles, clear descriptions, and quantified targets where relevant.\n\n"
        
            "5.2 Financial Restructuring\n"
            "  5.2.1 Debt Management\n"
            "    - Include specific recommendations for debt restructuring, cost reduction, and financing improvements.\n\n"
        
            "5.3 Market Development\n"
            "  5.3.1 Product Diversification\n"
            "    - Provide initiatives on new products, value-added formulations, sustainability products, and digital integration.\n"
            "  5.3.2 Market Expansion\n"
            "    - Include domestic market recovery strategy, new agriculture initiatives, export market development, and strategic partnerships.\n\n"
        
            "5.4 Governance and Sustainability\n"
            "  5.4.1 Corporate Governance Enhancement\n"
            "    - Recommendations to strengthen board composition, transparency, and management frameworks.\n"
            "  5.4.2 Environmental and Social Responsibility\n"
            "    - Initiatives for environmental performance, energy efficiency, community impact, and sustainable agriculture promotion.\n\n"
        
            "For each subheading, write in structured business report style with:\n"
            "- Bullet points where needed.\n"
            "- Bold the title of each recommendation.\n"
            "- Provide concise explanations (2-3 sentences).\n"
            "- Include quantified metrics or targets where applicable.\n"
            "- If external references or examples are used, provide real URLs in parentheses where possible.\n\n"
        
            "Ensure formatting is clear, organized, and professional for a strategic consulting report."
        ),
        temperature=0.4
    ).text.strip()
    doc.add_heading("5. Strategic Initiatives and Recommendations", level=1)
    doc.add_paragraph(strategic_initiatives_response)
    analysis_context['strategic_initiatives'] = strategic_initiatives_response

    context_summary = f"""
    EXECUTIVE SUMMARY:
    {analysis_context['executive_summary']}

    COMPANY OVERVIEW:   
    {analysis_context['company_overview']}

    INDUSTRY ANALYSIS:
    {analysis_context['industry_analysis']}

    FINANCIAL ANALYSIS:
    {analysis_context['financial_analysis']}

    SWOT ANALYSIS:
    {analysis_context['swot_analysis']}

    STRATEGIC INITIATIVES AND RECOMMENDATIONS:
    {analysis_context['strategic_initiatives']}
    """

    conclusion_response = co.chat(
        model="command-nightly",
        message=f"""
        You are a senior strategy consultant. Based on the following comprehensive company analysis, write Section 6: Conclusion.

        The company analysis is:

    {context_summary}

    Instructions:

    1. Summarize the company's current situation, key challenges, strategic opportunities, and recommended actions in a clear, professional tone.

    2. Categorize each final recommendation under one of the following: Exit, Monitor, Restructure, Invest, or Strategic Adjustments. Choose the most strategic fit.

    3. For each recommendation:
        - Start with the category and title in bold.
        - Provide a concise explanation (1-2 sentences).

    4. End with a strong final strategic recommendation sentence outlining the path forward for the company.

    Limit to a maximum of 200 words.
    Return in structured business report style with professional formatting.
    """,
        temperature=0.4
    ).text.strip()

# Add conclusion to report
    doc.add_heading("6. Conclusion", level=1)
    doc.add_paragraph(conclusion_response)

#################################################################


    output_dir = "data"
    os.makedirs(output_dir, exist_ok=True)
    filename = f"{custom_filename}.docx" if custom_filename else f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    # Generate Summary Overview Report
    overview_path = generate_summary_overview(
        analysis_context=analysis_context,
        company_name=company_name,
        industry_name=industry_name,
        output_dir=output_dir,
        custom_filename=overview_filename 
    )

    return output_path, overview_path


